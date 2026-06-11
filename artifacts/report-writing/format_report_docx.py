from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path("验收报告_ClockLink_终稿.docx")


def set_run_font(run, east_asia="宋体", ascii_font="Consolas", size=None, bold=None):
    run.font.name = ascii_font
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def main():
    document = Document(DOCX_PATH)

    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Consolas"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")

    for style_name, size in (
        ("Title", 20),
        ("Subtitle", 14),
        ("Heading 1", 16),
        ("Heading 2", 13),
        ("Heading 3", 11.5),
    ):
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Consolas"
            style.font.size = Pt(size)
            style.font.bold = True
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("【图 "):
            paragraph.style = styles["Normal"]
            paragraph.paragraph_format.left_indent = Cm(0.4)
            paragraph.paragraph_format.right_indent = Cm(0.4)
        for run in paragraph.runs:
            if text.startswith("【图 "):
                set_run_font(run, east_asia="宋体", ascii_font="Consolas", size=Pt(10), bold=False)

    for table in document.tables:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_border(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            east_asia="宋体",
                            ascii_font="Consolas",
                            size=Pt(9.5),
                            bold=True if row_index == 0 else None,
                        )

    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
